"""ingest_new_dataset.py
=============================================================================
Bring the live FAISS index up to date AND enrich every chunk with keywords.

Problem this fixes:
  - New_Dataset PDFs were chunked (chunks.jsonl=64k) but never embedded into the
    index (faiss.index=61k) -> chatbot can't find them.
  - New_Dataset .docx (19) and .xlsx (14) were never processed at all.
  - Chunks have no `keywords` field (wanted for thesis display).

What it does (safe full rebuild):
  1. Load existing chunks.jsonl.
  2. Extract + chunk new .docx (python-docx, incl. tables) and .xlsx (pandas, all
     sheets) from New_Dataset that aren't already in chunks.jsonl.
  3. Add a `keywords` field (top TF-IDF terms) to EVERY chunk.
  4. Embed all chunk texts with e5 ("passage: ", normalised) and build a fresh
     FAISS IndexFlatIP.
  5. Back up the old index/ dir, then write faiss.index + chunks_meta.json
     (now WITH keywords) + index_info.json.

USAGE
    python ingest_new_dataset.py --dry-run          # extract/chunk/keywords only, no embed
    python ingest_new_dataset.py --device cuda      # full rebuild on GPU
    python ingest_new_dataset.py --device cpu       # full rebuild on CPU (slow)

Run with the backend STOPPED and other apps closed (needs RAM + a few-GB index).
=============================================================================
"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import time
from pathlib import Path

import numpy as np

ROOT = Path(r"D:\Project\RAG_UPI")
NEW_DIR = ROOT / "Dataset" 
CHUNKS_JSONL = ROOT / "Dataset" / "_pipeline" / "chunked" / "chunks.jsonl"
INDEX_DIR = ROOT / "Dataset" / "_pipeline" / "index"
EMBED_MODEL = "intfloat/multilingual-e5-base"

CHUNK_SIZE = 1500
CHUNK_OVERLAP = 200
KEYWORDS_TOP = 8

# National research-funding NOISE (not UPI-specific): grant-recipient lists,
# BIMA / Kosabangsa / Panduan Penelitian. Dropped from the index when --upi-only.
NOISE_TITLE = re.compile(
    r"(daftar[_\s-]*(nama[_\s-]*)?penerima[_\s-]*pendanaan"
    r"|pengumuman[_\s-]*penerima[_\s-]*pendanaan"
    r"|lampiran.*penerima.*pendan"
    r"|panduan[_\s-]*penelitian([_\s-]*dan[_\s-]*pengabdian)?"
    r"|panduan.*pengabdian"
    r"|kosabangsa"
    r"|panduan[_\s-]*bima|^bima[_\s-]|[_\s-]bima[_\s-]"
    r")",
    re.IGNORECASE,
)
NOISE_TEXT = re.compile(
    r"(lldikti\s+wilayah|penerima\s+pendanaan|riset\s+kolaborasi\s+indonesia\s+\(rki\)"
    r"|skema\s+(penelitian|pengabdian)\s+(dasar|terapan))",
    re.IGNORECASE,
)


def is_noise(row: dict) -> bool:
    if NOISE_TITLE.search(row.get("title", "") or ""):
        return True
    if NOISE_TEXT.search(row.get("text", "") or ""):
        return True
    return False


def log(m: str) -> None:
    print(f"[ingest] {m}", flush=True)


_WS = re.compile(r"\s+")


def clean(t: str) -> str:
    return _WS.sub(" ", (t or "").replace(" ", " ")).strip()


def doc_id_for(path: Path) -> str:
    rel = str(path.relative_to(NEW_DIR)).lower()
    return hashlib.sha1(rel.encode("utf-8")).hexdigest()[:16]


# --------------------------------------------------------------------------- #
# Extraction
# --------------------------------------------------------------------------- #
def extract_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    parts: list[str] = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for tbl in d.tables:  # tables hold the real data in the FORMASI docs
        for row in tbl.rows:
            cells = [clean(c.text) for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx(path: Path) -> str:
    import pandas as pd
    sheets = pd.read_excel(path, sheet_name=None, header=None, dtype=str)
    out: list[str] = []
    for name, df in sheets.items():
        out.append(f"## Sheet: {name}")
        for _, row in df.iterrows():
            cells = [clean(str(c)) for c in row if str(c) != "nan" and clean(str(c))]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def chunk_text(text: str, size: int, overlap: int) -> list[str]:
    text = text.strip()
    if not text:
        return []
    chunks, i, n = [], 0, len(text)
    while i < n:
        end = min(i + size, n)
        # Snap the cut back to the nearest whitespace so a word / number / table
        # cell is never split in half (the "...Sistem Pembelajaran Da" problem).
        # Only backtrack within the last part of the window; if there is no space
        # there, fall back to the hard cut.
        if end < n:
            space = text.rfind(" ", i + int(size * 0.6), end)
            if space != -1:
                end = space
        chunks.append(text[i:end].strip())
        if end >= n:
            break
        i = max(end - overlap, i + 1)
    return [c for c in chunks if len(c) >= 40]


# --------------------------------------------------------------------------- #
def load_existing_chunks() -> list[dict]:
    rows = []
    with CHUNKS_JSONL.open(encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                rows.append(json.loads(line))
    return rows


def existing_doc_ids(rows: list[dict]) -> set[str]:
    return {r.get("doc_id") for r in rows}


def build_new_chunks(known_doc_ids: set[str]) -> list[dict]:
    """Extract + chunk New_Dataset .docx/.xlsx not already in the corpus."""
    new_rows: list[dict] = []
    files = sorted(list(NEW_DIR.rglob("*.docx")) + list(NEW_DIR.rglob("*.xlsx")))
    log(f"found {len(files)} docx/xlsx under New_Dataset")
    for path in files:
        if path.name.startswith("~$"):  # skip Office lock files
            continue
        did = doc_id_for(path)
        if did in known_doc_ids:
            continue
        category = path.relative_to(NEW_DIR).parts[0] if len(path.relative_to(NEW_DIR).parts) > 1 else "New_Dataset"
        try:
            text = extract_docx(path) if path.suffix.lower() == ".docx" else extract_xlsx(path)
        except Exception as exc:  # noqa: BLE001
            log(f"  ! gagal ekstrak {path.name}: {exc}")
            continue
        pieces = chunk_text(clean(text), CHUNK_SIZE, CHUNK_OVERLAP)
        for ci, piece in enumerate(pieces):
            new_rows.append({
                "doc_id": did,
                "source": str(path),
                "url": None,
                "title": path.stem,
                "category": category,
                "source_type": path.suffix.lower().lstrip("."),
                "page": 1,
                "section": None,
                "text": piece,
                "chunk_length": len(piece),
                "chunk_index": ci,
                "chunk_id": f"{did}::{ci}",
            })
        log(f"  + {path.name} [{category}] -> {len(pieces)} chunks")
    return new_rows


# --------------------------------------------------------------------------- #
# Keywords (TF-IDF top terms per chunk) — display metadata only
# --------------------------------------------------------------------------- #
ID_STOP = set("yang dan di ke dari untuk pada dengan atau adalah ini itu para "
              "sebagai dalam akan tidak juga oleh agar bagi serta dapat telah "
              "secara nomor tahun no the of and a an in on for to".split())


def add_keywords(rows: list[dict], top: int) -> None:
    from sklearn.feature_extraction.text import TfidfVectorizer
    log(f"computing TF-IDF keywords (top {top}/chunk) over {len(rows):,} chunks ...")
    vec = TfidfVectorizer(stop_words=list(ID_STOP), token_pattern=r"(?u)\b[a-zA-Z][a-zA-Z]{2,}\b",
                          max_features=50000, ngram_range=(1, 2), sublinear_tf=True)
    X = vec.fit_transform(clean(r.get("text", "")) for r in rows)
    terms = vec.get_feature_names_out()
    Xc = X.tocsr()
    for i, r in enumerate(rows):
        row = Xc.getrow(i)
        if row.nnz == 0:
            r["keywords"] = []
            continue
        idx = row.indices[np.argsort(row.data)[::-1][:top]]
        r["keywords"] = [terms[j] for j in idx]


# --------------------------------------------------------------------------- #
def main() -> None:
    ap = argparse.ArgumentParser(description="Rebuild index with New_Dataset docx/xlsx + keywords.")
    ap.add_argument("--dry-run", action="store_true", help="Extract/chunk/keywords only; no embed/index.")
    ap.add_argument("--device", choices=["cpu", "cuda", "auto"], default="auto")
    ap.add_argument("--keywords-top", type=int, default=KEYWORDS_TOP)
    ap.add_argument("--batch", type=int, default=16)
    ap.add_argument("--upi-only", action="store_true",
                    help="Exclude national research-funding noise docs (grant lists, "
                         "BIMA, Kosabangsa, Panduan Penelitian) from the INDEX. "
                         "chunks.jsonl keeps the full record; only the index is filtered.")
    args = ap.parse_args()

    t0 = time.time()
    existing = load_existing_chunks()
    log(f"existing chunks.jsonl: {len(existing):,}")
    known = existing_doc_ids(existing)

    new_rows = build_new_chunks(known)
    log(f"new docx/xlsx chunks: {len(new_rows):,}")

    all_rows = existing + new_rows
    log(f"total chunks to index: {len(all_rows):,}")

    add_keywords(all_rows, args.keywords_top)
    log("sample keywords: " + json.dumps(all_rows[0].get("keywords", []), ensure_ascii=False))

    # persist the merged chunks.jsonl (now includes new docs + keywords)
    if not args.dry_run:
        with CHUNKS_JSONL.open("w", encoding="utf-8") as f:
            for r in all_rows:
                f.write(json.dumps(r, ensure_ascii=False) + "\n")
        log(f"rewrote {CHUNKS_JSONL.name} with {len(all_rows):,} chunks (+keywords)")

    if args.dry_run:
        log(f"DRY-RUN done in {time.time()-t0:.1f}s — no embedding/index changes.")
        return

    # Index only UPI-relevant chunks when --upi-only (chunks.jsonl stays full).
    if args.upi_only:
        index_rows = [r for r in all_rows if not is_noise(r)]
        log(f"--upi-only: indexing {len(index_rows):,} chunks "
            f"({len(all_rows) - len(index_rows):,} national-noise chunks excluded)")
    else:
        index_rows = all_rows

    # --- embed + rebuild FAISS index ---
    import faiss
    import torch
    from sentence_transformers import SentenceTransformer
    device = args.device
    if device == "auto":
        device = "cuda" if torch.cuda.is_available() else "cpu"
    log(f"loading embedder {EMBED_MODEL} on {device} ...")
    model = SentenceTransformer(EMBED_MODEL, device=device)
    texts = ["passage: " + clean(r.get("text", "")) for r in index_rows]
    log(f"embedding {len(texts):,} chunks (this is the slow part) ...")
    vecs = model.encode(texts, batch_size=args.batch, convert_to_numpy=True,
                        normalize_embeddings=True, show_progress_bar=True).astype("float32")

    dim = vecs.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(vecs)
    log(f"FAISS index built: {index.ntotal:,} vectors, dim={dim}")

    # backup old index dir
    backup = INDEX_DIR.parent / f"index_backup_{time.strftime('%Y%m%d_%H%M%S')}"
    shutil.copytree(INDEX_DIR, backup)
    log(f"backed up old index -> {backup}")

    faiss.write_index(index, str(INDEX_DIR / "faiss.index"))
    (INDEX_DIR / "chunks_meta.json").write_text(
        json.dumps(index_rows, ensure_ascii=False), encoding="utf-8")
    (INDEX_DIR / "index_info.json").write_text(json.dumps({
        "embedding_model": EMBED_MODEL, "embedding_dim": int(dim),
        "n_vectors": int(index.ntotal), "use_e5_prefixes": True,
        "rebuilt_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "includes": "New_Dataset docx/xlsx + keywords",
    }, indent=2), encoding="utf-8")

    log("=" * 60)
    log(f"DONE in {time.time()-t0:.1f}s — index now {index.ntotal:,} vectors (was 61053).")
    log("Restart the backend to load the new index.")


if __name__ == "__main__":
    main()
