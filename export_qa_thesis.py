"""export_qa_thesis.py
=============================================================================
Turn the generated Q&A set (knowledge_layer/_rebuild/_work/*.jsonl) into
thesis-ready deliverables, neatly grouped per category:

  1. dataset_qa_skripsi.xlsx  - one sheet per category + a "Ringkasan" sheet,
                                full columns (id, pertanyaan, jawaban, dll).
  2. dataset_qa_skripsi.docx  - Lampiran-ready: summary table + one heading &
                                table per category (No | Pertanyaan | Jawaban
                                Singkat | Kesulitan). Drop straight into the
                                thesis appendix.

Reads the per-theme checkpoints so it also works on a partial (still-running)
generation. Categories are emitted in the fixed thesis order.

USAGE
    python export_qa_thesis.py
    python export_qa_thesis.py --src D:/.../_rebuild --out D:/.../_rebuild/exports
=============================================================================
"""
from __future__ import annotations

import argparse
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

RAG_ROOT = Path(r"D:\Project\RAG_UPI")
DEFAULT_SRC = RAG_ROOT / "knowledge_layer" / "_rebuild"

# fixed thesis order (label -> checkpoint code)
ORDER = [
    ("Penerimaan Mahasiswa Baru (PMB)",         "PMB"),
    ("Direktorat Pendidikan",                   "DITPEND"),
    ("LPPM",                                    "LPPM"),
    ("Informasi Publik (PPID)",                 "PPID"),
    ("UPI Kampus Cibiru",                       "CIBIRU"),
    ("UPI Kampus Sumedang",                     "SMD"),
    ("UPI Kampus Purwakarta",                   "PWK"),
    ("UPI Kampus Serang",                       "SRG"),
    ("UPI Kampus Tasikmalaya",                  "TSM"),
    ("Informasi Umum UPI & Kalender Akademik",  "UMUM"),
]

DIFF_ID = {"easy": "Mudah", "medium": "Sedang", "hard": "Sulit"}


def load_checkpoints(src: Path) -> dict[str, list[dict]]:
    work = src / "_work"
    by_cat: dict[str, list[dict]] = defaultdict(list)
    for label, code in ORDER:
        p = work / f"{code}.jsonl"
        if not p.exists():
            continue
        for line in p.open(encoding="utf-8"):
            if line.strip():
                by_cat[label].append(json.loads(line))
    return by_cat


# --------------------------------------------------------------------------- #
# Excel
# --------------------------------------------------------------------------- #
def build_xlsx(by_cat: dict[str, list[dict]], out: Path) -> Path:
    wb = openpyxl.Workbook()
    thin = Side(style="thin", color="BBBBBB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    head_fill = PatternFill("solid", fgColor="1F4E79")
    head_font = Font(bold=True, color="FFFFFF")
    wrap = Alignment(wrap_text=True, vertical="top")
    center = Alignment(horizontal="center", vertical="center")

    # summary sheet
    ws = wb.active
    ws.title = "Ringkasan"
    ws.append(["No", "Kategori", "Jumlah Pertanyaan"])
    for c in range(1, 4):
        cell = ws.cell(row=1, column=c)
        cell.fill = head_fill
        cell.font = head_font
        cell.border = border
        cell.alignment = center
    total = 0
    for i, (label, _) in enumerate(ORDER, 1):
        n = len(by_cat.get(label, []))
        total += n
        ws.append([i, label, n])
        for c in range(1, 4):
            ws.cell(row=i + 1, column=c).border = border
    ws.append(["", "TOTAL", total])
    r = ws.max_row
    ws.cell(row=r, column=2).font = Font(bold=True)
    ws.cell(row=r, column=3).font = Font(bold=True)
    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 42
    ws.column_dimensions["C"].width = 20

    cols = ["No", "ID", "Pertanyaan", "Jawaban Singkat", "Jawaban Lengkap",
            "Kesulitan", "Persona", "Dokumen Sumber", "Chunk Sumber"]
    widths = [5, 12, 45, 45, 60, 11, 16, 20, 24]
    for label, code in ORDER:
        rows = by_cat.get(label, [])
        if not rows:
            continue
        title = code[:31]  # sheet name limit
        ws = wb.create_sheet(title=title)
        ws.append(cols)
        for c in range(1, len(cols) + 1):
            cell = ws.cell(row=1, column=c)
            cell.fill = head_fill
            cell.font = head_font
            cell.border = border
            cell.alignment = center
        for i, rec in enumerate(rows, 1):
            ws.append([
                i,
                rec.get("question_id", ""),
                rec.get("question", ""),
                rec.get("answer_short", ""),
                rec.get("answer_long", ""),
                DIFF_ID.get(rec.get("difficulty", "medium"), "Sedang"),
                rec.get("persona", ""),
                ", ".join(rec.get("supporting_documents", []) or []),
                ", ".join(rec.get("supporting_chunks", []) or []),
            ])
            for c in range(1, len(cols) + 1):
                cell = ws.cell(row=i + 1, column=c)
                cell.border = border
                cell.alignment = wrap
        for idx, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(idx)].width = w
        ws.freeze_panes = "A2"
    wb.save(out)
    return out


# --------------------------------------------------------------------------- #
# Word (Lampiran-ready)
# --------------------------------------------------------------------------- #
def _shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_widths(table, widths_inch):
    from docx.shared import Inches
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_inch):
            row.cells[i].width = Inches(w)


def build_docx(by_cat: dict[str, list[dict]], out: Path) -> Path:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)

    h = doc.add_heading("LAMPIRAN: DATASET PERTANYAAN-JAWABAN EVALUASI", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Dataset ini berisi pasangan pertanyaan dan jawaban acuan (reference "
        "answer) yang di-generate secara otomatis dan tervalidasi dari korpus "
        "dokumen resmi Universitas Pendidikan Indonesia. Setiap pertanyaan "
        "terhubung ke dokumen sumbernya (grounded), dikelompokkan berdasarkan "
        "tema/departemen, dan digunakan sebagai basis pengujian retrieval serta "
        "evaluasi RAGAS.")

    # summary table
    doc.add_heading("Ringkasan Distribusi Pertanyaan", level=2)
    total = sum(len(by_cat.get(l, [])) for l, _ in ORDER)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "No", "Kategori", "Jumlah"
    for c in hdr:
        _shade(c, "1F4E79")
        for p in c.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (label, _) in enumerate(ORDER, 1):
        cells = tbl.add_row().cells
        cells[0].text = str(i)
        cells[1].text = label
        cells[2].text = str(len(by_cat.get(label, [])))
    trow = tbl.add_row().cells
    trow[0].text = ""
    trow[1].text = "TOTAL"
    trow[2].text = str(total)
    for c in (trow[1], trow[2]):
        for p in c.paragraphs:
            for run in p.runs:
                run.font.bold = True
    _set_widths(tbl, [0.5, 4.2, 1.0])

    # per-category tables
    for label, code in ORDER:
        rows = by_cat.get(label, [])
        if not rows:
            continue
        doc.add_page_break()
        doc.add_heading(f"{label}  (n = {len(rows)})", level=2)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "No"
        hdr[1].text = "Pertanyaan"
        hdr[2].text = "Jawaban Singkat"
        hdr[3].text = "Kesulitan"
        for c in hdr:
            _shade(c, "2E75B6")
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for i, rec in enumerate(rows, 1):
            cells = tbl.add_row().cells
            cells[0].text = str(i)
            cells[1].text = rec.get("question", "")
            cells[2].text = rec.get("answer_short", "")
            cells[3].text = DIFF_ID.get(rec.get("difficulty", "medium"), "Sedang")
        _set_widths(tbl, [0.4, 3.1, 2.8, 0.8])

    doc.save(out)
    return out


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", default=str(DEFAULT_SRC))
    ap.add_argument("--out", default="")
    args = ap.parse_args()

    src = Path(args.src)
    out_dir = Path(args.out) if args.out else (src / "exports")
    out_dir.mkdir(parents=True, exist_ok=True)

    by_cat = load_checkpoints(src)
    total = sum(len(v) for v in by_cat.values())
    print(f"[export] loaded {total} questions across {len(by_cat)} categories")
    for label, _ in ORDER:
        print(f"[export]   {len(by_cat.get(label, [])):5d}  {label}")

    xlsx = build_xlsx(by_cat, out_dir / "dataset_qa_skripsi.xlsx")
    print(f"[export] wrote {xlsx}")
    docx = build_docx(by_cat, out_dir / "dataset_qa_skripsi.docx")
    print(f"[export] wrote {docx}")
    print("[export] done")


if __name__ == "__main__":
    main()
