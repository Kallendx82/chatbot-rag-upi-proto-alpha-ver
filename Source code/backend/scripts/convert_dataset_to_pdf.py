#!/usr/bin/env python3
"""
Dataset preprocessing: Convert all non-PDF files (Excel, Word, images, etc.)
to PDF format so they can be served via the /api/source/{doc_id}.pdf endpoint.

Usage:
    python convert_dataset_to_pdf.py <dataset_root_dir>

Example:
    python convert_dataset_to_pdf.py D:\\Project\\RAG_UPI\\Dataset

Dependencies (install via pip):
    pip install openpyxl python-docx pillow reportlab PyPDF2
"""
import sys
from pathlib import Path
from typing import Optional
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)
log = logging.getLogger(__name__)


def convert_excel_to_pdf(xlsx_path: Path, pdf_path: Path) -> bool:
    """Convert Excel file to PDF using openpyxl + reportlab."""
    try:
        from openpyxl import load_workbook
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors

        wb = load_workbook(str(xlsx_path))
        ws = wb.active

        # Extract table data
        data = []
        for row in ws.iter_rows(values_only=True):
            data.append(list(row))

        # Create PDF
        doc = SimpleDocTemplate(str(pdf_path), pagesize=A4)
        elements = []

        if data:
            table = Table(data, colWidths=[None]*len(data[0]))
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            elements.append(table)

        doc.build(elements)
        log.info(f"✓ Converted Excel → PDF: {xlsx_path.name}")
        return True
    except Exception as e:
        log.error(f"✗ Failed to convert Excel {xlsx_path.name}: {e}")
        return False


def convert_word_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """Convert Word document to PDF using python-docx + reportlab."""
    try:
        from docx import Document as DocxDocument
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        doc = DocxDocument(str(docx_path))
        styles = getSampleStyleSheet()

        pdf_doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
        elements = []

        for para in doc.paragraphs:
            if para.text.strip():
                elements.append(Paragraph(para.text, styles['Normal']))
                elements.append(Spacer(1, 12))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        elements.append(Paragraph(cell.text, styles['Normal']))

        pdf_doc.build(elements)
        log.info(f"✓ Converted Word → PDF: {docx_path.name}")
        return True
    except Exception as e:
        log.error(f"✗ Failed to convert Word {docx_path.name}: {e}")
        return False


def convert_image_to_pdf(img_path: Path, pdf_path: Path) -> bool:
    """Convert image file to PDF."""
    try:
        from PIL import Image

        img = Image.open(str(img_path))

        # Convert RGBA to RGB if needed
        if img.mode in ("RGBA", "LA", "P"):
            bg = Image.new("RGB", img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[-1] if img.mode == "RGBA" else None)
            img = bg

        img.save(str(pdf_path), "PDF")
        log.info(f"✓ Converted Image → PDF: {img_path.name}")
        return True
    except Exception as e:
        log.error(f"✗ Failed to convert Image {img_path.name}: {e}")
        return False


def scan_and_convert(dataset_root: Path) -> dict:
    """Scan dataset and convert all non-PDF files to PDF."""
    stats = {
        "total": 0,
        "xlsx": 0,
        "docx": 0,
        "doc": 0,
        "images": 0,
        "other": 0,
        "converted": 0,
        "failed": 0,
    }

    # File type handlers
    handlers = {
        ".xlsx": convert_excel_to_pdf,
        ".xls": convert_excel_to_pdf,
        ".docx": convert_word_to_pdf,
        ".doc": convert_word_to_pdf,
        ".png": convert_image_to_pdf,
        ".jpg": convert_image_to_pdf,
        ".jpeg": convert_image_to_pdf,
        ".gif": convert_image_to_pdf,
        ".bmp": convert_image_to_pdf,
    }

    log.info(f"Scanning dataset at: {dataset_root}")

    # Find all non-PDF files
    for file_path in dataset_root.rglob("*"):
        if not file_path.is_file():
            continue

        suffix = file_path.suffix.lower()
        stats["total"] += 1

        # Skip PDFs
        if suffix == ".pdf":
            continue

        # Track file types
        if suffix == ".xlsx":
            stats["xlsx"] += 1
        elif suffix in (".doc", ".docx"):
            stats["docx"] += 1 if suffix == ".docx" else 0
            stats["doc"] += 1 if suffix == ".doc" else 0
        elif suffix in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
            stats["images"] += 1
        else:
            stats["other"] += 1
            continue  # Skip unknown types

        # Create PDF version
        pdf_path = file_path.with_suffix(".pdf")

        # Skip if PDF already exists
        if pdf_path.exists():
            log.info(f"⊘ PDF already exists, skipping: {pdf_path.name}")
            continue

        # Convert file
        handler = handlers.get(suffix)
        if handler:
            if handler(file_path, pdf_path):
                stats["converted"] += 1
            else:
                stats["failed"] += 1

    return stats


def main():
    if len(sys.argv) < 2:
        print("Usage: python convert_dataset_to_pdf.py <dataset_root_dir>")
        print("Example: python convert_dataset_to_pdf.py D:\\Project\\RAG_UPI\\Dataset")
        sys.exit(1)

    dataset_root = Path(sys.argv[1])

    if not dataset_root.exists():
        log.error(f"Dataset directory not found: {dataset_root}")
        sys.exit(1)

    log.info("=" * 60)
    log.info("Dataset PDF Conversion Tool")
    log.info("=" * 60)

    stats = scan_and_convert(dataset_root)

    log.info("=" * 60)
    log.info("Conversion Summary:")
    log.info(f"  Total files: {stats['total']}")
    log.info(f"  Excel files: {stats['xlsx']}")
    log.info(f"  Word files: {stats['docx'] + stats['doc']}")
    log.info(f"  Image files: {stats['images']}")
    log.info(f"  Other types: {stats['other']}")
    log.info(f"  Converted: {stats['converted']}")
    log.info(f"  Failed: {stats['failed']}")
    log.info("=" * 60)

    if stats["failed"] > 0:
        log.warning(f"⚠ {stats['failed']} file(s) failed to convert. Check logs above.")
        sys.exit(1)

    log.info("✓ Conversion complete!")
    sys.exit(0)


if __name__ == "__main__":
    main()
